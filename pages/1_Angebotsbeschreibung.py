"""
Teil 1 — Angebotsbeschreibung
==============================
Streamlit page for teachers to complete the Weiterbildungsmodulbeschreibung.
Pre-filled from the Weiterbildungs-Radar session state where possible.
Outputs: downloadable Word (.docx) and optional Nextcloud upload.
"""
import io, json, math, re, datetime
from pathlib import Path

import streamlit as st
import pandas as pd

st.set_page_config(
    page_title="Angebotsbeschreibung · TH Wildau",
    page_icon=None,
    layout="wide",
)

DATA = Path(__file__).parent.parent / "data"

# ─── HELPERS ─────────────────────────────────────────────────────────

def get(key, default=""):
    """Read from session_state, fall back to default."""
    return st.session_state.get(key, default) or default

def load_lernziel_templates():
    path = DATA / "lernziel_templates.json"
    if path.exists():
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return {"default": []}

def section(colour, title):
    st.markdown(
        f'<div style="background:{colour};border-radius:8px;padding:8px 16px;'
        f'margin:1.5rem 0 0.8rem 0"><span style="font-size:1rem;font-weight:600">'
        f'{title}</span></div>',
        unsafe_allow_html=True,
    )

def hint(text):
    st.caption(f"💡 {text}")

# ─── PRE-FILL FROM RADAR ────────────────────────────────────────────

# Radar passes data via session_state.radar_params and session_state.confirmed_berufe
radar = st.session_state.get("radar_params", {})
confirmed_berufe = st.session_state.get("confirmed_berufe", set())

# DQR level derived from degree
DEGREE_TO_DQR = {
    "Bachelor":                          "DQR Niveau 6 — Bachelor-Ebene",
    "Zertifikat / Hochschulzertifikat":  "DQR Niveau 6 — Bachelor-Ebene",
    "Microcredential (digitales Badge / Teilleistung)": "DQR Niveau 6 — Bachelor-Ebene",
    "Teilnahmebescheinigung":            "DQR Niveau 6 — Bachelor-Ebene",
    "Master":                            "DQR Niveau 7 — Master-Ebene",
    "Abschlussprüfung":                  "DQR Niveau 7 — Master-Ebene",
}
dqr_prefill = DEGREE_TO_DQR.get(radar.get("degree",""), "DQR Niveau 6 — Bachelor-Ebene")

# Workload calculation: ECTS × 25-30 hours
ects_prefill = radar.get("ects", 0) or 0
workload_min = ects_prefill * 25
workload_max = ects_prefill * 30

# Lernziel suggestions based on selected categories
templates = load_lernziel_templates()
selected_cats = radar.get("selected_cats", [])
suggested_lernziele = []
for cat in selected_cats:
    suggested_lernziele.extend(templates.get(cat, []))
if not suggested_lernziele:
    suggested_lernziele = templates.get("default", [])

# ─── PAGE ────────────────────────────────────────────────────────────

st.title("Teil 1 — Angebotsbeschreibung")
st.markdown(
    "Dieses Formular entspricht dem offiziellen TH-Wildau-Formular zur Konzeption von "
    "Weiterbildungsangeboten. Felder die automatisch aus dem Weiterbildungs-Radar "
    "übertragen wurden sind mit **★** markiert. Bitte prüfen und ergänzen Sie alle Angaben."
)
st.markdown("---")

# ═══════════════════════════════════════════════════════════════════
# BLOCK 1: Grunddaten
# ═══════════════════════════════════════════════════════════════════
section("#dceefb", "Grunddaten des Angebots")

col1, col2 = st.columns(2)

with col1:
    name = st.text_input(
        "★ Name des Weiterbildungsangebots",
        value=radar.get("title",""),
        help="Aus dem Radar übertragen. Bitte ggf. anpassen.",
    )

    abschluss = st.selectbox(
        "★ Abschluss",
        ["Zertifikat / Hochschulzertifikat",
         "Microcredential (digitales Badge / Teilleistung)",
         "Teilnahmebescheinigung",
         "Bachelor",
         "Master",
         "Certificate of Advanced Studies (CAS)",
         "Aufbauzertifikat",
         "Abschlussprüfung",
         "Sonstiges"],
        index=["Zertifikat / Hochschulzertifikat",
               "Microcredential (digitales Badge / Teilleistung)",
               "Teilnahmebescheinigung","Bachelor","Master",
               "Certificate of Advanced Studies (CAS)",
               "Aufbauzertifikat","Abschlussprüfung","Sonstiges"
               ].index(radar.get("degree","Zertifikat / Hochschulzertifikat"))
        if radar.get("degree") in ["Zertifikat / Hochschulzertifikat",
               "Microcredential (digitales Badge / Teilleistung)",
               "Teilnahmebescheinigung","Bachelor","Master",
               "Certificate of Advanced Studies (CAS)",
               "Aufbauzertifikat","Abschlussprüfung","Sonstiges"] else 0,
    )

    fachbereich = st.radio(
        "Zuordnung Fachbereich",
        ["INW", "WIR", "Nicht eindeutig zuzuordnen"],
        horizontal=True,
    )

with col2:
    wiss_leitung = st.text_input(
        "Wissenschaftliche Leitung / Verantwortlich für die Konzeption",
        placeholder="Name, Titel, Professur",
    )
    verantwortlich_durchfuehrung = st.text_input(
        "Verantwortlich für die Durchführung",
        placeholder="Name oder Stelle, falls abweichend",
    )
    kooperationspartner = st.text_input(
        "Kooperationspartner (falls zutreffend)",
        placeholder="z.B. IHK Cottbus, Siemens AG",
    )

# ═══════════════════════════════════════════════════════════════════
# BLOCK 2: Inhalte und Lernziele
# ═══════════════════════════════════════════════════════════════════
section("#d4edda", "Inhalte und Lernziele")

st.markdown("**★ Kurzbeschreibung** (aus dem Radar übertragen — bitte ergänzen)")
inhalte_text = st.text_area(
    "Inhalte und Lernziele",
    value=radar.get("description",""),
    height=120,
    label_visibility="collapsed",
)

if suggested_lernziele:
    st.markdown("**Vorgeschlagene Lernziele** — anklicken zum Hinzufügen:")
    hint("Diese Vorschläge basieren auf den thematischen Schwerpunkten aus dem Radar.")

    if "selected_lernziele" not in st.session_state:
        st.session_state.selected_lernziele = set()

    cols = st.columns(2)
    for i, lz in enumerate(suggested_lernziele):
        is_sel = lz in st.session_state.selected_lernziele
        bg = "#0f6e56" if is_sel else "#f0f7f4"
        col = "#fff" if is_sel else "#085041"
        with cols[i % 2]:
            st.markdown(
                f'<div style="background:{bg};color:{col};border-radius:6px;'
                f'padding:6px 12px;margin:2px 0;font-size:13px">{lz}</div>',
                unsafe_allow_html=True,
            )
            if st.button(
                "× Entfernen" if is_sel else "+ Hinzufügen",
                key=f"lz_{i}",
                use_container_width=True,
            ):
                if is_sel:
                    st.session_state.selected_lernziele.discard(lz)
                else:
                    st.session_state.selected_lernziele.add(lz)
                st.rerun()

    if st.session_state.selected_lernziele:
        st.markdown("**Ausgewählte Lernziele:**")
        for lz in sorted(st.session_state.selected_lernziele):
            st.markdown(f"- {lz}")

# ═══════════════════════════════════════════════════════════════════
# BLOCK 3: Niveau, Umfang, Format
# ═══════════════════════════════════════════════════════════════════
section("#fff0d4", "Niveau, Umfang und Format")

c1, c2, c3 = st.columns(3)

with c1:
    dqr = st.selectbox(
        "★ Niveau (DQR)",
        ["DQR Niveau 6 — Bachelor-Ebene", "DQR Niveau 7 — Master-Ebene"],
        index=0 if "6" in dqr_prefill else 1,
    )
    hint(f"Automatisch abgeleitet aus Abschluss: {radar.get('degree','—')}")

with c2:
    ects = st.number_input(
        "★ Anzahl CP (ECTS)",
        min_value=0, max_value=120,
        value=int(ects_prefill),
    )
    if ects > 0:
        st.caption(f"Entspricht {ects*25}–{ects*30} Zeitstunden Workload")

with c3:
    workload_display = f"{ects*25}–{ects*30}" if ects > 0 else "—"
    st.text_input(
        "★ Zeitlicher Aufwand (Zeitstunden)",
        value=workload_display,
        disabled=True,
        help="Automatisch berechnet: CP × 25–30 Stunden",
    )

c4, c5, c6 = st.columns(3)

with c4:
    FORMAT_MAP = {
        "Online / Digital":    "Online / Digital",
        "Hybrid / Blended":    "Hybrid / Blended",
        "Präsenz":             "Präsenz",
        "Noch offen":          "Noch offen",
    }
    format_opts = ["Online / Digital","Hybrid / Blended","Präsenz",
                   "Synchron","Asynchron","Synchron und asynchron","Noch offen"]
    radar_fmt = radar.get("format","Noch offen")
    fmt_idx = next((i for i,f in enumerate(format_opts) if radar_fmt in f), 6)
    teilnahme_form = st.selectbox(
        "★ Form der Teilnahme",
        format_opts,
        index=fmt_idx,
    )

with c5:
    monate = radar.get("months", 0) or 0
    dauer_str = f"{monate} Monate" if monate else ""
    dauer = st.text_input(
        "★ Gesamtdauer des Angebots",
        value=dauer_str,
        placeholder="z.B. 6 Monate, 2 Semester, 5 Tage",
    )

with c6:
    turnus = st.selectbox(
        "Angebotsturnus",
        ["Jedes Sommersemester","Jedes Wintersemester","Jedes Semester",
         "Jährlich","Quartalsweise","Unregelmäßig","Nach Bedarf"],
    )
    turnus_custom = st.text_input("Angebotsturnus — Freifeld (optional)", placeholder="z.B. erstmalig WiSe 2026/27")

# ═══════════════════════════════════════════════════════════════════
# BLOCK 4: Zielgruppe und Voraussetzungen
# ═══════════════════════════════════════════════════════════════════
section("#fce4ec", "Zielgruppe und Voraussetzungen")

# Zielgruppe from confirmed Berufe
berufe_list = sorted(confirmed_berufe) if confirmed_berufe else []
if berufe_list:
    st.markdown("**★ Zielgruppe — aus dem Radar übertragen:**")
    hint("Diese Berufsgruppen wurden im Radar als Zielgruppe identifiziert. Ergänzen oder entfernen Sie Einträge.")
    zielgruppe_berufe = st.multiselect(
        "Zielgruppen-Berufsbilder",
        options=berufe_list,
        default=berufe_list,
        label_visibility="collapsed",
    )
else:
    st.info("Keine Berufsgruppen aus dem Radar — bitte manuell eingeben.")
    zielgruppe_berufe = []

zielgruppe_freitext = st.text_area(
    "Zielgruppe (Beschreibung / Ergänzung)",
    placeholder="z.B. Fach- und Führungskräfte mit mind. 2 Jahren Berufserfahrung im Bereich …",
    height=80,
)

st.markdown("**Voraussetzungen für die Teilnahme**")
voraus_std = st.multiselect(
    "Standardformulierungen (Mehrfachauswahl möglich)",
    [
        "Abgeschlossenes Hochschulstudium (Bachelor oder gleichwertig)",
        "Abgeschlossenes Hochschulstudium (Master oder gleichwertig)",
        "Einschlägige Berufserfahrung von mindestens 1 Jahr",
        "Einschlägige Berufserfahrung von mindestens 3 Jahren",
        "Keine formalen Voraussetzungen",
        "Berufsausbildung in einem einschlägigen Fachgebiet",
        "Grundkenntnisse in [Thema] werden empfohlen",
    ],
)
voraus_custom = st.text_area(
    "Weitere Voraussetzungen (Freitext)",
    placeholder="Besondere Regelungen, empfohlene Vorkenntnisse, technische Anforderungen …",
    height=60,
)

# ═══════════════════════════════════════════════════════════════════
# BLOCK 5: Curriculare Struktur
# ═══════════════════════════════════════════════════════════════════
section("#ede0f5", "Curriculare Struktur des Angebots")

struktur_beschr = st.text_area(
    "Struktur (Kurzbeschreibung)",
    placeholder="z.B. 1 Pflichtmodul (15 CP) + 3 Wahlpflichtmodule (je 5 CP). "
                "Module können unabhängig voneinander belegt werden.",
    height=80,
)

st.markdown("**Pflichtmodule**")
hint("Fügen Sie Module hinzu. CP pro Modul können eingegeben werden.")

if "pflicht_module" not in st.session_state:
    st.session_state.pflicht_module = [{"modul":"","cp":0,"herkunft":"Weiterbildungsmodul","pruefung":"","teilnahme":""}]

for i, mod in enumerate(st.session_state.pflicht_module):
    c1,c2,c3,c4,c5,c6 = st.columns([4,1,3,3,3,1])
    mod["modul"]    = c1.text_input("Modul", value=mod["modul"], key=f"pm_{i}_name", label_visibility="collapsed" if i>0 else "visible")
    mod["cp"]       = c2.number_input("CP", value=mod["cp"], min_value=0, max_value=30, key=f"pm_{i}_cp", label_visibility="collapsed" if i>0 else "visible")
    mod["herkunft"] = c3.selectbox("Herkunft", ["Weiterbildungsmodul","TH-Wildau-Studiengang","Kooperationspartner","Sonstiges"], key=f"pm_{i}_her", label_visibility="collapsed" if i>0 else "visible")
    mod["pruefung"] = c4.selectbox("Prüfungsform", ["Klausur","Hausarbeit","Projektarbeit","Präsentation","Portfolio","Referat","Mündliche Prüfung","Praktische Arbeit","Keine Prüfung"], key=f"pm_{i}_prf", label_visibility="collapsed" if i>0 else "visible")
    mod["teilnahme"]= c5.selectbox("Teilnahmeform", ["Präsenz","Online","Hybrid","Blended"], key=f"pm_{i}_tn", label_visibility="collapsed" if i>0 else "visible")
    if c6.button("×", key=f"pm_{i}_del") and len(st.session_state.pflicht_module) > 1:
        st.session_state.pflicht_module.pop(i); st.rerun()

if st.button("+ Pflichtmodul hinzufügen"):
    st.session_state.pflicht_module.append({"modul":"","cp":0,"herkunft":"Weiterbildungsmodul","pruefung":"","teilnahme":""})
    st.rerun()

st.markdown("**Wahlpflichtmodule**")
if "wahl_module" not in st.session_state:
    st.session_state.wahl_module = []

for i, mod in enumerate(st.session_state.wahl_module):
    c1,c2,c3,c4,c5,c6 = st.columns([4,1,3,3,3,1])
    mod["modul"]    = c1.text_input("Modul", value=mod["modul"], key=f"wm_{i}_name", label_visibility="collapsed")
    mod["cp"]       = c2.number_input("CP", value=mod["cp"], min_value=0, max_value=30, key=f"wm_{i}_cp", label_visibility="collapsed")
    mod["herkunft"] = c3.selectbox("Herkunft", ["Weiterbildungsmodul","TH-Wildau-Studiengang","Kooperationspartner","Sonstiges"], key=f"wm_{i}_her", label_visibility="collapsed")
    mod["pruefung"] = c4.selectbox("Prüfungsform", ["Klausur","Hausarbeit","Projektarbeit","Präsentation","Portfolio","Referat","Mündliche Prüfung","Praktische Arbeit","Keine Prüfung"], key=f"wm_{i}_prf", label_visibility="collapsed")
    mod["teilnahme"]= c5.selectbox("Teilnahmeform", ["Präsenz","Online","Hybrid","Blended"], key=f"wm_{i}_tn", label_visibility="collapsed")
    if c6.button("×", key=f"wm_{i}_del"):
        st.session_state.wahl_module.pop(i); st.rerun()

if st.button("+ Wahlpflichtmodul hinzufügen"):
    st.session_state.wahl_module.append({"modul":"","cp":0,"herkunft":"Weiterbildungsmodul","pruefung":"","teilnahme":""})
    st.rerun()

# Total CP check
total_cp = (sum(m["cp"] for m in st.session_state.pflicht_module) +
            sum(m["cp"] for m in st.session_state.wahl_module))
if total_cp > 0 and ects > 0:
    if total_cp == ects:
        st.success(f"CP-Summe: {total_cp} CP — stimmt mit den ECTS überein.")
    else:
        st.warning(f"CP-Summe der Module: {total_cp} CP — ECTS-Angabe oben: {ects} CP. Bitte prüfen.")

# ═══════════════════════════════════════════════════════════════════
# BLOCK 6: Prüfungen (summary)
# ═══════════════════════════════════════════════════════════════════
section("#f0f4ff", "Prüfungen (Übersicht)")
st.caption("Automatisch aus den Modulen oben zusammengefasst — kann manuell ergänzt werden.")
pruef_auto = ", ".join(
    set(m["pruefung"] for m in st.session_state.pflicht_module + st.session_state.wahl_module
        if m["pruefung"] and m["pruefung"] != "Keine Prüfung")
) or "—"
pruefungen = st.text_area("Prüfungsformen", value=pruef_auto, height=60)

# ═══════════════════════════════════════════════════════════════════
# GENERATE WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════════

def collect_form_data():
    """Bundle all form values into one dict for document generation."""
    return {
        "name": name,
        "abschluss": abschluss,
        "fachbereich": fachbereich,
        "wiss_leitung": wiss_leitung,
        "verantwortlich_durchfuehrung": verantwortlich_durchfuehrung,
        "kooperationspartner": kooperationspartner,
        "inhalte": inhalte_text,
        "lernziele_selected": sorted(st.session_state.get("selected_lernziele", set())),
        "dqr": dqr,
        "ects": ects,
        "workload": f"{ects*25}–{ects*30} Stunden" if ects > 0 else "—",
        "teilnahme_form": teilnahme_form,
        "dauer": dauer,
        "turnus": f"{turnus}" + (f" / {turnus_custom}" if turnus_custom else ""),
        "zielgruppe_berufe": zielgruppe_berufe,
        "zielgruppe_freitext": zielgruppe_freitext,
        "voraussetzungen": voraus_std + ([voraus_custom] if voraus_custom else []),
        "struktur": struktur_beschr,
        "pflicht_module": st.session_state.pflicht_module,
        "wahl_module": st.session_state.wahl_module,
        "pruefungen": pruefungen,
        "datum": datetime.date.today().strftime("%d.%m.%Y"),
    }

def build_docx(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section_obj in doc.sections:
        section_obj.top_margin    = Cm(2)
        section_obj.bottom_margin = Cm(2)
        section_obj.left_margin   = Cm(2.5)
        section_obj.right_margin  = Cm(2.5)

    def add_heading(text, level=1, colour="1a3a6b"):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(colour)
        return p

    def add_field(label, value, prefilled=False):
        p = doc.add_paragraph()
        run_label = p.add_run(label + ": ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        marker = "★ " if prefilled else ""
        run_val = p.add_run(marker + (str(value) if value else "—"))
        run_val.font.size = Pt(10)
        if prefilled:
            run_val.font.color.rgb = RGBColor(0, 80, 60)
        p.paragraph_format.space_after = Pt(4)

    def add_table_row(table, cells):
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            row.cells[i].text = str(cell_text)
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        return row

    # ── Title ──
    title_p = doc.add_paragraph()
    title_r = title_p.add_run("Formular zur Konzeption von Weiterbildungsangeboten")
    title_r.bold = True; title_r.font.size = Pt(16)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run(f"Teil 1 — Angebotsbeschreibung  ·  Stand: {data['datum']}")
    sub_r.font.size = Pt(10); sub_r.font.color.rgb = RGBColor(100,100,100)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Block 1: Grunddaten ──
    add_heading("1. Grunddaten des Angebots", 1)
    add_field("Name des Weiterbildungsangebots", data["name"], prefilled=True)
    add_field("Abschluss", data["abschluss"], prefilled=True)
    add_field("Zuordnung Fachbereich", data["fachbereich"])
    add_field("Wissenschaftliche Leitung", data["wiss_leitung"])
    add_field("Verantwortlich für die Durchführung", data["verantwortlich_durchfuehrung"])
    add_field("Kooperationspartner", data["kooperationspartner"] or "Nicht zutreffend")

    # ── Block 2: Inhalte ──
    add_heading("2. Inhalte und Lernziele", 1)
    p = doc.add_paragraph()
    p.add_run("Kurzbeschreibung:").bold = True
    doc.add_paragraph(data["inhalte"] or "—").paragraph_format.space_after = Pt(6)

    if data["lernziele_selected"]:
        p2 = doc.add_paragraph()
        p2.add_run("Lernziele:").bold = True
        for lz in data["lernziele_selected"]:
            doc.add_paragraph(lz, style="List Bullet")

    # ── Block 3: Niveau / Umfang ──
    add_heading("3. Niveau, Umfang und Format", 1)
    add_field("Niveau (DQR)", data["dqr"], prefilled=True)
    add_field("CP (ECTS)", str(data["ects"]), prefilled=True)
    add_field("Zeitlicher Aufwand", data["workload"], prefilled=True)
    add_field("Form der Teilnahme", data["teilnahme_form"], prefilled=True)
    add_field("Gesamtdauer", data["dauer"], prefilled=True)
    add_field("Angebotsturnus", data["turnus"])

    # ── Block 4: Zielgruppe ──
    add_heading("4. Zielgruppe und Voraussetzungen", 1)
    if data["zielgruppe_berufe"]:
        p = doc.add_paragraph()
        p.add_run("Zielgruppen-Berufsbilder (★ aus Radar):").bold = True
        for beruf in data["zielgruppe_berufe"]:
            doc.add_paragraph(beruf, style="List Bullet")

    if data["zielgruppe_freitext"]:
        add_field("Zielgruppe (Beschreibung)", data["zielgruppe_freitext"])

    if data["voraussetzungen"]:
        p = doc.add_paragraph()
        p.add_run("Voraussetzungen für die Teilnahme:").bold = True
        for v in data["voraussetzungen"]:
            if v:
                doc.add_paragraph(v, style="List Bullet")

    # ── Block 5: Curriculare Struktur ──
    add_heading("5. Curriculare Struktur", 1)
    if data["struktur"]:
        doc.add_paragraph(data["struktur"])

    def add_module_table(title, modules):
        if not modules: return
        doc.add_paragraph().add_run(title).bold = True
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Modul","CP","Herkunft","Prüfungsform","Teilnahmeform"]):
            hdr[i].text = h
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.bold = True; run.font.size = Pt(9)
        for mod in modules:
            if mod["modul"]:
                add_table_row(tbl, [
                    mod["modul"], str(mod["cp"]), mod["herkunft"],
                    mod["pruefung"], mod["teilnahme"]
                ])
        doc.add_paragraph()

    add_module_table("Pflichtmodule", data["pflicht_module"])
    add_module_table("Wahlpflichtmodule", data["wahl_module"])

    # ── Block 6: Prüfungen ──
    add_heading("6. Prüfungen", 1)
    doc.add_paragraph(data["pruefungen"] or "—")

    # ── Footer note ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("★ = automatisch aus dem Weiterbildungs-Radar der TH Wildau übertragen")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(100,100,100)
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ═══════════════════════════════════════════════════════════════════
# NEXTCLOUD UPLOAD
# ═══════════════════════════════════════════════════════════════════

def upload_to_nextcloud(docx_bytes: bytes, filename: str,
                         nc_url: str, nc_user: str, nc_password: str,
                         nc_path: str) -> tuple[bool, str]:
    try:
        import requests
        url = f"{nc_url.rstrip('/')}/remote.php/dav/files/{nc_user}/{nc_path.strip('/')}/{filename}"
        resp = requests.put(
            url,
            data=docx_bytes,
            auth=(nc_user, nc_password),
            headers={"Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
            timeout=30,
        )
        if resp.status_code in (200, 201, 204):
            return True, url
        return False, f"HTTP {resp.status_code}: {resp.text[:200]}"
    except Exception as e:
        return False, str(e)

# ═══════════════════════════════════════════════════════════════════
# OUTPUT SECTION
# ═══════════════════════════════════════════════════════════════════
st.markdown("---")
section("#f5f5f5", "Dokument erstellen und speichern")

col_dl, col_nc = st.columns(2)

with col_dl:
    st.markdown("**Als Word-Dokument herunterladen**")
    if st.button("Angebotsbeschreibung als .docx erstellen", type="primary"):
        with st.spinner("Dokument wird erstellt..."):
            data = collect_form_data()
            docx_bytes = build_docx(data)
            safe_name = re.sub(r'[^\w\-]', '_', name or "Angebotsbeschreibung")
            filename = f"Angebotsbeschreibung_{safe_name}_{datetime.date.today()}.docx"
        st.download_button(
            label="Download starten",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

with col_nc:
    st.markdown("**In Nextcloud speichern**")
    with st.expander("Nextcloud-Verbindung einrichten"):
        nc_url  = st.text_input("Nextcloud-URL", placeholder="https://cloud.th-wildau.de", key="nc_url")
        nc_user = st.text_input("Benutzername", key="nc_user")
        nc_pass = st.text_input("Passwort / App-Passwort", type="password", key="nc_pass")
        nc_path = st.text_input("Zielordner in Nextcloud",
                                 value="Weiterbildung/Angebotsbeschreibungen",
                                 key="nc_path")
        st.caption("Tipp: Verwenden Sie ein App-Passwort (Nextcloud → Einstellungen → Sicherheit).")

    if st.button("In Nextcloud hochladen"):
        if not all([nc_url, nc_user, nc_pass]):
            st.error("Bitte Nextcloud-URL, Benutzername und Passwort eingeben.")
        else:
            with st.spinner("Wird hochgeladen..."):
                data = collect_form_data()
                docx_bytes = build_docx(data)
                safe_name = re.sub(r'[^\w\-]', '_', name or "Angebotsbeschreibung")
                filename = f"Angebotsbeschreibung_{safe_name}_{datetime.date.today()}.docx"
                ok, msg = upload_to_nextcloud(docx_bytes, filename, nc_url, nc_user, nc_pass, nc_path)
            if ok:
                st.success(f"Erfolgreich hochgeladen: {nc_path}/{filename}")
            else:
                st.error(f"Fehler beim Upload: {msg}")

# ── Back link ──
st.markdown("---")
if st.button("← Zurück zum Weiterbildungs-Radar"):
    st.switch_page("app.py")
